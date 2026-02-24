# -*- coding: utf-8 -*-
"""
生成测试Word文档
用于测试VBA导入工具
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    """创建测试Word文档"""
    
    # 创建文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('学生信息管理系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph('VBA测试文档')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分隔线
    doc.add_paragraph('_' * 50)
    
    # 添加使用说明
    doc.add_heading('使用说明', level=1)
    
    instructions = [
        '1. 本文档包含VBA代码，可用于学生登录和信息管理',
        '2. 按 Alt + F11 可查看和编辑VBA代码',
        '3. 运行宏：按 Alt + F8 选择"显示登录窗体"',
        '4. 使用本工具可导入/导出VBA代码',
    ]
    
    for inst in instructions:
        doc.add_paragraph(inst)
    
    # 添加学生信息表格
    doc.add_heading('学生信息表', level=1)
    
    # 创建表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '学号'
    header_cells[1].text = '姓名'
    header_cells[2].text = '记录时间'
    
    # 添加示例数据
    data = [
        ('2024001', '张三', '2024-01-01 10:00:00'),
        ('2024002', '李四', '2024-01-01 11:00:00'),
        ('2024003', '王五', '2024-01-01 12:00:00'),
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        for j, value in enumerate(row_data):
            row_cells[j].text = value
    
    # 添加页脚
    doc.add_paragraph('_' * 50)
    footer = doc.add_paragraph('请在Word中按 Alt+F11 查看VBA代码')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    output_path = r'E:\pydemo\03VBA工程\demo\测试文档.docx'
    doc.save(output_path)
    print(f"测试文档已创建: {output_path}")
    
    return output_path


if __name__ == '__main__':
    create_test_document()
